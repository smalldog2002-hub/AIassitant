import streamlit as st
import google.generativeai as genai
from PIL import Image
import time
import os
import tempfile
import docx
import ebooklib
from ebooklib import epub
from bs4 import BeautifulSoup
from gtts import gTTS
from io import BytesIO

# --- 页面全局配置 ---
st.set_page_config(
    page_title="汪汪的视觉全能助手",
    page_icon="🔮",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- 核心 CSS 美化 ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Noto+Sans+SC:wght@400;700;900&display=swap');
    
    .stApp {
        background: linear-gradient(135deg, #f5f7fa 0%, #c3cfe2 100%);
        font-family: 'Noto Sans SC', sans-serif;
    }

    .main-header {
        font-size: 3rem;
        font-weight: 900;
        background: linear-gradient(to right, #4f46e5, #06b6d4);
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        text-align: center;
        margin-bottom: 0.5rem;
        letter-spacing: -1px;
    }
    
    .sub-header {
        text-align: center;
        color: #64748b;
        font-size: 1rem;
        margin-bottom: 2.5rem;
        font-weight: 500;
    }

    .glass-card {
        background: rgba(255, 255, 255, 0.7);
        backdrop-filter: blur(10px);
        -webkit-backdrop-filter: blur(10px);
        border-radius: 24px;
        border: 1px solid rgba(255, 255, 255, 0.5);
        box-shadow: 0 8px 32px 0 rgba(31, 38, 135, 0.05);
        padding: 2rem;
        margin-bottom: 2rem;
    }

    .stButton>button {
        width: 100%;
        border-radius: 50px;
        height: 3.5rem;
        font-weight: 700;
        border: none;
        background: linear-gradient(90deg, #4f46e5 0%, #6366f1 100%);
        color: white;
        box-shadow: 0 4px 15px rgba(79, 102, 241, 0.3);
        transition: all 0.3s ease;
    }
    .stButton>button:hover {
        transform: scale(1.02);
        box-shadow: 0 6px 20px rgba(79, 102, 241, 0.4);
        color: white !important;
    }

    .ai-output-box {
        background-color: #ffffff;
        border-left: 6px solid #4f46e5;
        padding: 2rem;
        border-radius: 0 16px 16px 0;
        box-shadow: 0 2px 10px rgba(0,0,0,0.02);
        margin-top: 1.5rem;
        line-height: 1.7;
    }
    
    /* 聊天气泡优化 */
    .chat-container { display: flex; flex-direction: column; gap: 15px; margin-bottom: 20px; }
    .chat-bubble { padding: 15px 20px; border-radius: 18px; max-width: 85%; font-size: 1rem; line-height: 1.5; box-shadow: 0 2px 5px rgba(0,0,0,0.03); }
    .chat-user { align-self: flex-end; background: linear-gradient(135deg, #6366f1 0%, #8b5cf6 100%); color: white; border-bottom-right-radius: 4px; }
    .chat-ai { align-self: flex-start; background: white; color: #1e293b; border-bottom-left-radius: 4px; border: 1px solid #f1f5f9; }
    
    .warning-box { background-color: #fef2f2; border: 1px solid #fee2e2; color: #991b1b; padding: 1rem; border-radius: 12px; display: flex; align-items: center; gap: 10px; font-size: 0.9rem; }
    
    /* 口语修正框 */
    .correction-box { background-color: #ecfdf5; border: 1px solid #a7f3d0; color: #047857; padding: 10px; border-radius: 8px; font-size: 0.9rem; margin-top: 5px; }
</style>
""", unsafe_allow_html=True)

# --- 侧边栏配置 ---
with st.sidebar:
    st.title("🔮 神经中枢")
    try:
        secrets_key = st.secrets.get("GEMINI_API_KEY", "")
    except FileNotFoundError:
        secrets_key = ""

    if secrets_key:
        st.success("✅ 视觉神经已连接 (Secrets)")
        if st.toggle("🔧 切换手动 Key"):
            api_key = st.text_input("输入新 Key", type="password")
        else:
            api_key = secrets_key
    else:
        api_key = st.text_input("🔑 API Key", type="password", help="在此输入 Key 激活所有功能")
    
    st.markdown("---")
    selected_mode = st.radio(
        "启用功能模块",
        [
            "🗣️ 口语陪练教练", # 新增模块
            "📸 你拍我答 (万能问答)",
            "💬 一起聊天吧 (全知全能)",
            "📚 全库文档问答 (PDF/Word/Epub)",
            "⚖️ 法律合同审查 (Word/PDF)",
            "🎙️ 会议纪要生成器",
            "🏥 医疗健康助手",
            "💻 自动化脚本写手",
            "✨ 社交配文生成"
        ]
    )
    st.caption("🚀 Core: gyuniku 1.5/2.5 Flash")

# --- 核心逻辑函数 ---

def get_model():
    if not api_key:
        st.error("🛑 神经中枢未连接：请配置 API Key")
        return None
    genai.configure(api_key=api_key)
    return genai.GenerativeModel('gemini-2.5-flash-preview-09-2025')

def generate_speech(text, lang_code):
    """生成语音 (gTTS)"""
    try:
        if not text: return None
        # gTTS 语言代码映射
        tts_lang = lang_code
        if lang_code == 'ko-KR': tts_lang = 'ko'
        elif lang_code == 'ja-JP': tts_lang = 'ja'
        elif lang_code == 'en-US': tts_lang = 'en'
        elif lang_code == 'fr-FR': tts_lang = 'fr'
        elif lang_code == 'th-TH': tts_lang = 'th'
        
        tts = gTTS(text=text, lang=tts_lang)
        fp = BytesIO()
        tts.write_to_fp(fp)
        return fp.getvalue()
    except Exception: return None

def extract_text_from_docx(file_path):
    try:
        doc = docx.Document(file_path)
        full_text = [para.text for para in doc.paragraphs if para.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Word 解析错误: {e}")
        return ""

def extract_text_from_epub(file_path):
    try:
        book = epub.read_epub(file_path)
        full_text = []
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                soup = BeautifulSoup(item.get_content(), 'html.parser')
                full_text.append(soup.get_text())
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Epub 解析错误: {e}")
        return ""

def process_and_upload(uploaded_file):
    file_ext = os.path.splitext(uploaded_file.name)[1].lower()
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_ext) as tmp_src:
        tmp_src.write(uploaded_file.getvalue())
        tmp_src_path = tmp_src.name

    final_path = tmp_src_path
    mime_type = "" 

    try:
        with st.status(f"📡 正在处理 {file_ext} 文件...", expanded=True) as status:
            if file_ext in ['.pdf', '.jpg', '.jpeg', '.png', '.webp', '.mp3', '.wav', '.aiff', '.aac', '.ogg', '.flac']:
                st.write(f"🚀 检测到原生支持格式，正在直传云端...")
                if file_ext == '.pdf': mime_type = 'application/pdf'
                elif file_ext in ['.jpg', '.jpeg']: mime_type = 'image/jpeg'
                elif file_ext == '.png': mime_type = 'image/png'
            elif file_ext in ['.docx', '.epub', '.txt', '.md', '.py', '.js', '.c', '.json']:
                st.write(f"🔄 正在解析 {file_ext} 文档结构...")
                text_content = ""
                if file_ext == '.docx': text_content = extract_text_from_docx(tmp_src_path)
                elif file_ext == '.epub': text_content = extract_text_from_epub(tmp_src_path)
                else:
                    with open(tmp_src_path, "r", encoding="utf-8", errors='ignore') as f:
                        text_content = f.read()
                
                if not text_content.strip(): raise ValueError(f"文档为空。")
                
                with tempfile.NamedTemporaryFile(delete=False, suffix=".txt", mode="w", encoding="utf-8") as tmp_txt:
                    tmp_txt.write(text_content)
                    final_path = tmp_txt.name
                mime_type = "text/plain"
                st.write("✅ 文档已转换为 AI 可读文本流")
            else:
                raise ValueError(f"暂不支持的文件格式: {file_ext}")

            st.write("☁️ 正在上传至 AI 知识库...")
            if mime_type: myfile = genai.upload_file(final_path, mime_type=mime_type)
            else: myfile = genai.upload_file(final_path)
            
            st.write("🧠 AI 正在构建上下文索引...")
            while myfile.state.name == "PROCESSING":
                time.sleep(1)
                myfile = genai.get_file(myfile.name)
            
            if myfile.state.name == "FAILED":
                status.update(label="❌ 文件处理失败", state="error")
                raise ValueError("Gemini 无法处理此文件")
            
            status.update(label="✅ 文件已挂载到 AI 大脑", state="complete")
            return myfile

    finally:
        if os.path.exists(tmp_src_path): os.remove(tmp_src_path)
        if final_path != tmp_src_path and os.path.exists(final_path): os.remove(final_path)

def render_ai_response(response_text):
    st.markdown(f"""<div class="ai-output-box">{response_text}</div>""", unsafe_allow_html=True)

# --- 主界面 ---

st.markdown('<div class="main-header">AI 视觉全能助手</div>', unsafe_allow_html=True)
st.markdown(f'<div class="sub-header">当前激活模块：<span style="color:#4f46e5; font-weight:bold;">{selected_mode}</span></div>', unsafe_allow_html=True)

# 0. 口语陪练教练 (新增)
if "口语" in selected_mode:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    
    # 状态管理
    if "practice_history" not in st.session_state: st.session_state.practice_history = []
    
    c1, c2, c3 = st.columns(3)
    with c1:
        target_lang = st.selectbox("🎯 目标语言", ["韩语 (Korean)", "英语 (English)", "日语 (Japanese)", "法语 (French)", "泰语 (Thai)"])
    with c2:
        scenario = st.selectbox("🎬 练习场景", ["日常闲聊", "餐厅点餐", "旅行问路", "初次见面", "商务会议"])
    with c3:
        st.write("")
        st.write("")
        if st.button("🔄 重置对话"):
            st.session_state.practice_history = []
            st.rerun()
            
    # 获取语言代码
    lang_code_map = {"韩语 (Korean)": "ko-KR", "英语 (English)": "en-US", "日语 (Japanese)": "ja-JP", "法语 (French)": "fr-FR", "泰语 (Thai)": "th-TH"}
    lang_code = lang_code_map[target_lang]
    
    # 初始化开场白
    if not st.session_state.practice_history:
        model = get_model()
        if model:
            init_prompt = f"你现在是一位地道的{target_lang}母语者。请用{target_lang}向我打招呼，并发起一个关于'{scenario}'的话题。请只输出{target_lang}，不要带翻译。"
            try:
                res = model.generate_content(init_prompt)
                st.session_state.practice_history.append({"role": "assistant", "text": res.text, "audio": None})
            except: pass

    # 显示聊天记录
    for msg in st.session_state.practice_history:
        role = msg["role"]
        text = msg["text"]
        css = "chat-ai" if role == "assistant" else "chat-user"
        
        st.markdown(f'<div class="chat-container"><div class="chat-bubble {css}">{text}</div></div>', unsafe_allow_html=True)
        
        # 只有 AI 的回复才生成语音按钮
        if role == "assistant":
            # 如果还没有音频缓存，尝试生成
            if msg.get("audio") is None:
                audio_data = generate_speech(text, lang_code)
                msg["audio"] = audio_data # 缓存起来
            
            if msg.get("audio"):
                st.audio(msg["audio"], format="audio/mp3", start_time=0)
            
            # 显示修正建议 (如果有)
            if "correction" in msg and msg["correction"]:
                st.markdown(f'<div class="correction-box">💡 <strong>语法建议：</strong> {msg["correction"]}</div>', unsafe_allow_html=True)

    # 输入框
    user_input = st.chat_input(f"用{target_lang}回复...")
    
    if user_input:
        st.session_state.practice_history.append({"role": "user", "text": user_input})
        st.rerun()

    # 处理 AI 回复
    if st.session_state.practice_history and st.session_state.practice_history[-1]["role"] == "user":
        last_input = st.session_state.practice_history[-1]["text"]
        
        with st.spinner("AI 老师正在思考..."):
            model = get_model()
            if model:
                try:
                    # 复杂的 Prompt：既要回复，又要纠错
                    prompt = f"""
                    你是一位{target_lang}口语老师。用户刚刚说了："{last_input}"。
                    当前场景：{scenario}。
                    
                    任务：
                    1. 像真人一样用{target_lang}自然地回复用户，继续对话。
                    2. 检查用户的输入是否有严重的语法错误或不自然的表达。
                    
                    请以 JSON 格式输出：
                    {{
                        "reply": "你的回复内容(仅{target_lang})",
                        "correction": "如果用户有错，用中文简短指出并给出正确说法；如果没错，留空字符串"
                    }}
                    """
                    
                    response = model.generate_content(prompt)
                    import json
                    try:
                        # 尝试解析 JSON
                        clean_json = response.text.strip()
                        if "```json" in clean_json:
                            clean_json = clean_json.split("```json")[1].split("```")[0]
                        data = json.loads(clean_json)
                        reply_text = data.get("reply", "")
                        correction = data.get("correction", "")
                    except:
                        # 兜底：如果没按 JSON 输出，直接用文本
                        reply_text = response.text
                        correction = ""
                    
                    st.session_state.practice_history.append({
                        "role": "assistant", 
                        "text": reply_text, 
                        "correction": correction,
                        "audio": None # 待生成
                    })
                    st.rerun()
                except Exception as e: st.error(f"Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

# 0.5 你拍我答
elif "你拍我答" in selected_mode:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.info("💡 解题、识物、翻译。支持图片、PDF。")
    tab1, tab2 = st.tabs(["📂 上传文件", "📸 拍照"])
    with tab1: file_up = st.file_uploader("支持 JPG, PNG, PDF", type=['jpg','png','jpeg', 'pdf'])
    with tab2: cam_up = st.camera_input("直接拍摄")
    target = file_up if file_up else cam_up
    
    if target:
        if hasattr(target, 'type') and 'pdf' in target.type:
            st.markdown(f"📄 **PDF 已就绪**: `{target.name}`")
        else:
            st.image(target, width=300)
        
        user_q = st.text_area("✍️ 请输入问题 (留空则默认解读)", height=80)
        if st.button("🚀 开始解答", type="primary"):
            model = get_model()
            if model:
                try:
                    gemini_file = process_and_upload(target)
                    q_prompt = user_q if user_q else "请详细解读这份内容。"
                    with st.spinner("🧠 AI 正在思考..."):
                        response = model.generate_content([q_prompt, gemini_file])
                        render_ai_response(response.text)
                except Exception as e: st.error(f"Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

# 0.8 聊天
elif "聊天" in selected_mode:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    if "general_chat_history" not in st.session_state: st.session_state.general_chat_history = []
    
    for role, text in st.session_state.general_chat_history:
        css_class = "chat-user" if role == "user" else "chat-ai"
        st.markdown(f'<div class="chat-container"><div class="chat-bubble {css_class}">{text}</div></div>', unsafe_allow_html=True)
        
    if query := st.chat_input("和我聊聊吧..."):
        st.session_state.general_chat_history.append(("user", query))
        st.rerun()

    if st.session_state.general_chat_history and st.session_state.general_chat_history[-1][0] == "user":
        with st.spinner("AI 正在思考..."):
            model = get_model()
            if model:
                try:
                    history_text = "\n".join([f"{r}: {t}" for r, t in st.session_state.general_chat_history[-10:]])
                    system_prompt = "你是一位全知全能、幽默风趣的 AI 助手。严禁讨论色情暴力话题。"
                    full_prompt = f"{system_prompt}\n\n历史：\n{history_text}\n\nAI 回复："
                    response = model.generate_content(full_prompt)
                    st.session_state.general_chat_history.append(("assistant", response.text))
                    st.rerun()
                except Exception as e: st.error(f"回复失败: {e}")
    
    if st.button("🗑️ 清空记录"):
        st.session_state.general_chat_history = []
        st.rerun()
    st.markdown('</div>', unsafe_allow_html=True)

# 2 & 5. 全库问答 + 合同审查
elif "全库" in selected_mode or "合同" in selected_mode:
    is_chat = "全库" in selected_mode
    
    if "doc_history" not in st.session_state: st.session_state.doc_history = []
    if "current_doc" not in st.session_state: st.session_state.current_doc = None
    if "current_name" not in st.session_state: st.session_state.current_name = None

    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    col1, col2 = st.columns([3, 1])
    with col1:
        supported_types = ['pdf', 'docx', 'epub', 'txt', 'md']
        if not is_chat: supported_types.extend(['jpg', 'png', 'jpeg']) 
        label_text = "📂 上传文档 (支持 PDF, Word .docx, Epub, Txt)"
        uploaded_doc = st.file_uploader(label_text, type=supported_types)
    with col2:
        st.write("") 
        st.write("") 
        if st.button("🔄 清空历史"):
            st.session_state.doc_history = []
            st.rerun()

    if uploaded_doc:
        if st.session_state.current_name != uploaded_doc.name:
            model = get_model()
            if model:
                try:
                    gemini_file = process_and_upload(uploaded_doc)
                    st.session_state.current_doc = gemini_file
                    st.session_state.current_name = uploaded_doc.name
                    st.session_state.doc_history = []
                except Exception as e: st.error(f"Load Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

    if st.session_state.current_doc:
        if not is_chat: # 合同审查
            if st.button("⚡ 开始深度风险审查", type="primary"):
                model = get_model()
                with st.spinner("⚖️ AI 法务正在审阅..."):
                    prompt = """
                    你是一位资深法律顾问。请严格审查这份合同文件。
                    输出一份《法律风险评估报告》，包含：高风险条款预警、权益保障缺失、修改建议(表格)、总体评分。
                    """
                    try:
                        response = model.generate_content([prompt, st.session_state.current_doc])
                        render_ai_response(response.text)
                    except Exception as e: st.error(f"Analysis Error: {e}")
        
        else: # 全库问答
            st.markdown("### 💬 知识库对话")
            for role, text in st.session_state.doc_history:
                css_class = "chat-user" if role == "user" else "chat-ai"
                st.markdown(f'<div class="chat-container"><div class="chat-bubble {css_class}">{text}</div></div>', unsafe_allow_html=True)
            
            if query := st.chat_input("关于这份文档，你想知道什么？"):
                st.session_state.doc_history.append(("user", query))
                st.rerun()
                
            if st.session_state.doc_history and st.session_state.doc_history[-1][0] == "user":
                last_query = st.session_state.doc_history[-1][1]
                if len(st.session_state.doc_history) % 2 != 0:
                    with st.spinner("AI 正在阅读..."):
                        model = get_model()
                        try:
                            response = model.generate_content([st.session_state.current_doc, last_query])
                            st.session_state.doc_history.append(("assistant", response.text))
                            st.rerun()
                        except Exception as e: st.error(f"Chat Error: {e}")

# 自动化脚本
elif "脚本" in selected_mode:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.info("💡 描述需求，AI 将为你编写 Python 自动化脚本。")
    script_requirement = st.text_area("需求描述...", height=150)
    if script_requirement and st.button("⚡ 生成代码", type="primary"):
        model = get_model()
        if model:
            with st.spinner("编写中..."):
                try:
                    response = model.generate_content(f"写一个Python脚本：{script_requirement}。要求：健壮、有注释。")
                    render_ai_response(response.text)
                except Exception as e: st.error(f"Error: {e}")
    st.markdown('</div>', unsafe_allow_html=True)

# 医疗 / 会议 / 其他
else:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    
    # 区分模式
    if "会议" in selected_mode:
        st.info("💡 支持 mp3, wav, m4a, ogg 等音频格式。")
        up_label = "上传音频"
        up_types = ['mp3', 'wav', 'm4a', 'ogg', 'flac']
        prompt_template = "请作为专业的首席会议秘书，根据录音生成一份完美的会议纪要。包含摘要、决策清单、待办事项和详细回顾。"
    elif "卡路里" in selected_mode:
        st.info("🍎 AI 营养师准备就绪")
        up_label = "上传食物图"
        up_types = ['jpg','png','jpeg']
        prompt_template = "分析食物，列出热量/营养成分表及建议。"
    elif "手写" in selected_mode:
        st.info("📝 OCR 识别引擎准备就绪")
        up_label = "上传笔记图"
        up_types = ['jpg','png','jpeg']
        prompt_template = "OCR 识别，转为电子文本，保留格式。"
    elif "配文" in selected_mode:
        st.info("✨ 创意文案引擎准备就绪")
        up_label = "上传图片"
        up_types = ['jpg','png','jpeg']
        pass 
    elif "医疗" in selected_mode:
         st.info("🏥 AI 医疗助手准备就绪")
         up_label = "上传报告/药盒"
         up_types = ['jpg','png','pdf']
         pass 

    # 统一上传逻辑
    if "配文" not in selected_mode and "医疗" not in selected_mode:
        col1, col2 = st.tabs(["📂 上传文件", "📸 拍照"])
        with col1: up_file = st.file_uploader(up_label, type=up_types)
        with col2: cam_file = st.camera_input("拍照")
        
        target = up_file if up_file else cam_file
        
        if target:
            if "mp3" not in getattr(target, 'type', '') and "wav" not in getattr(target, 'type', ''):
                st.image(target, width=400)
            
            if st.button("开始分析", type="primary"):
                model = get_model()
                if model:
                    with st.spinner("分析中..."):
                        try:
                            g_file = process_and_upload(target)
                            response = model.generate_content([prompt_template, g_file])
                            render_ai_response(response.text)
                        except Exception as e: st.error(f"Error: {e}")

    # 配文
    elif "配文" in selected_mode:
        col1, col2 = st.tabs(["📂 上传图片", "📸 拍照"])
        with col1: up_file = st.file_uploader("上传图片", type=['jpg','png','jpeg'])
        with col2: cam_file = st.camera_input("拍照")
        target = up_file if up_file else cam_file
        
        if target:
            st.image(target, width=300)
            style = st.selectbox("文案风格", ["文艺清新", "幽默搞笑", "扎心语录", "小红书爆款"])
            if st.button("✨ 生成文案", type="primary"):
                model = get_model()
                if model:
                    with st.spinner("创作中..."):
                        try:
                            g_file = Image.open(target)
                            response = model.generate_content([f"写3条{style}风格的朋友圈文案，带Emoji。", g_file])
                            render_ai_response(response.text)
                        except Exception as e: st.error(f"Error: {e}")

    # 医疗
    elif "医疗" in selected_mode:
        med_type = st.radio("任务", ["体检解读", "药品识别"], horizontal=True)
        col1, col2 = st.tabs(["📂 上传", "📸 拍照"])
        with col1: up_file = st.file_uploader("文件", type=['jpg','png','pdf'])
        with col2: cam_file = st.camera_input("拍照")
        target = up_file if up_file else cam_file
        
        if target:
            if st.button("开始分析", type="primary"):
                model = get_model()
                if model:
                    with st.spinner("诊断中..."):
                        try:
                            g_file = process_and_upload(target)
                            prompt = "解读体检报告" if "体检" in med_type else "解读药品说明书"
                            response = model.generate_content([prompt, g_file])
                            render_ai_response(response.text)
                            st.markdown("""<div class="warning-box">⚠️ 结果仅供参考，不作为医疗依据。</div>""", unsafe_allow_html=True)
                        except Exception as e: st.error(f"Error: {e}")

# --- 页脚 ---
st.markdown("---")

st.markdown('<div style="text-align: center; color: #94a3b8; font-size: 0.8rem;">Powered by <strong>gyuniku 养乐多益力多 多多益善 1.5/2.5 Flash Vision</strong> | Built with Streamlit</div>', unsafe_allow_html=True)

